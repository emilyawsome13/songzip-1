from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(93, 175, 207)
TEXT = RGBColor(22, 36, 49)
TEXT_SOFT = RGBColor(82, 101, 117)
LINE = "DCE6EE"
BOX_FILL = "EDF7FA"
BOX_EDGE = "BFDDE6"


POLICIES = [
    {
        "title": "SongZip Privacy Policy",
        "subtitle": "Current dashboard build transparency draft",
        "summary": (
            "SongZip is intended for personal research, evaluation, and "
            "entertainment-related organization workflows. Where account linking is "
            "offered, it uses official provider authorization, requests limited "
            "permissions, and is not intended to sell personal data, use borrowed "
            "credentials, or hide how provider data is accessed."
        ),
        "output": "privacy-policy.docx",
        "sections": [
            {
                "title": "1. Scope and purpose",
                "paragraphs": [
                    "SongZip is positioned as a personal workflow dashboard for queueing media-related requests, checking progress, and optionally linking a user's own provider account through official OAuth flows when a feature needs it.",
                    "This policy explains what limited data SongZip may process, why it may process it, and how the service should avoid misleading users about privacy or account access.",
                ],
            },
            {
                "title": "2. Data SongZip may process",
                "paragraphs": [
                    "If account linking is enabled, SongZip may need to process a limited set of account and authorization data so the requested feature can work.",
                    "If OAuth is enabled, SongZip should not claim that it collects 'no user data.' The more accurate statement is that it is intended to process only the limited authorization and account data needed for the user-requested feature.",
                ],
                "bullets": [
                    "Provider account identifiers",
                    "Display names or email-style account labels returned by the provider",
                    "Access tokens and refresh tokens",
                    "Granted scopes and token expiry timestamps",
                    "Session identifiers and connection-status metadata",
                ],
            },
            {
                "title": "3. How data should be used",
                "bullets": [
                    "To complete official account-linking flows initiated by the user",
                    "To maintain the user's active session and connected-account status",
                    "To refresh tokens when needed for the exact user-facing feature that was authorized",
                    "To provide visible dashboard features tied to the linked account",
                ],
                "paragraphs": [
                    "SongZip is not intended to use provider or user data for unrelated profiling, advertising, or resale.",
                ],
            },
            {
                "title": "4. Data use limits",
                "bullets": [
                    "SongZip should not sell or rent personal data.",
                    "SongZip should not use personal data for ad targeting or unrelated profiling.",
                    "SongZip should not ask users to paste raw tokens into the interface.",
                    "SongZip should not accept borrowed, shared, leaked, or third-party credentials.",
                    "SongZip should not use linked-account data to claim it can bypass platform quotas or reviews.",
                ],
            },
            {
                "title": "5. Permissions and provider access",
                "paragraphs": [
                    "SongZip should use official documented access methods and request only the minimum relevant permissions for the feature being provided.",
                ],
                "bullets": [
                    "Users should connect their own Spotify or Google accounts through official provider sign-in pages.",
                    "Permissions should be requested in context and tied to a visible feature.",
                    "If a new feature requires broader permissions, disclosures should be updated before access is requested.",
                ],
            },
            {
                "title": "6. Retention, security, and current limitations",
                "paragraphs": [
                    "Before public deployment, SongZip should implement encrypted secret storage, clear retention rules, disconnect and deletion workflows, and access controls around any stored connection records.",
                    "The current local prototype stores limited OAuth connection records server-side for operational use. Until encryption and retention controls are hardened, it should not be described as a production-grade privacy implementation.",
                ],
            },
            {
                "title": "7. Spotify and Google / YouTube commitments",
                "bullets": [
                    "Provide accurate privacy disclosures for Spotify data access and use.",
                    "Avoid misleading users about how provider data is accessed or what permissions are needed.",
                    "Respect minimum-permission expectations in Spotify and Google policies.",
                    "Avoid prohibited uses of provider data, including unauthorized transfers or deceptive access patterns.",
                ],
            },
        ],
    },
    {
        "title": "SongZip Acceptable Use Policy",
        "subtitle": "Current dashboard build transparency draft",
        "summary": (
            "SongZip is intended for personal research, evaluation, and "
            "entertainment-related organization workflows. It is not intended for "
            "unauthorized redistribution, borrowed credentials, quota-circumvention "
            "claims, or deceptive use of provider APIs."
        ),
        "output": "acceptable-use-policy.docx",
        "sections": [
            {
                "title": "1. Intended use",
                "paragraphs": [
                    "SongZip is intended for limited personal workflows tied to user-requested media organization and evaluation.",
                ],
                "bullets": [
                    "Personal research and evaluation of media metadata or workflow behavior",
                    "Entertainment-related organization of user-requested media references",
                    "User-initiated account connections through official provider authorization flows",
                ],
            },
            {
                "title": "2. Prohibited use",
                "bullets": [
                    "Unauthorized copying, redistribution, resale, or public distribution of protected content",
                    "Using borrowed, shared, leaked, or third-party tokens, credentials, or example accounts",
                    "Evading rate limits, quotas, access controls, or platform review requirements",
                    "Scraping provider data outside official documented access methods",
                    "Misleading users or platforms about what the product does, how it accesses data, or what it stores",
                ],
            },
            {
                "title": "3. Account and token rules",
                "bullets": [
                    "Users should connect only their own accounts through official OAuth flows.",
                    "SongZip should not ask users to paste raw access tokens or refresh tokens into the interface.",
                    "SongZip should not represent linked user accounts as a way to transfer app-level quota responsibility away from the operator.",
                ],
            },
            {
                "title": "4. Quota and rate-limit position",
                "paragraphs": [
                    "Linking a user's own Spotify or Google account does not automatically remove app-wide developer obligations around rate limits, quotas, audits, or platform review.",
                ],
                "bullets": [
                    "SongZip should not advertise quota evasion as a feature.",
                    "SongZip should not encourage users to rotate borrowed accounts or credentials.",
                    "SongZip should use caching, deduplication, and normal rate-limit handling instead of circumvention.",
                ],
            },
            {
                "title": "5. Content and copyright position",
                "paragraphs": [
                    "SongZip does not claim ownership of third-party content made available through Spotify, YouTube, Google, or other providers.",
                    "Users remain responsible for ensuring that their use of any content is lawful and permitted by provider terms, copyright law, and any attached subscription or license conditions.",
                    "If a provider's terms prohibit a particular use, that use should be treated as prohibited within SongZip as well.",
                ],
            },
            {
                "title": "6. Public-positioning guardrails",
                "paragraphs": [
                    "If SongZip is offered with paid subscriptions, broad consumer marketing, or features that copy or redistribute protected content, describing it publicly as 'strictly for research only' may be misleading unless the real product truly matches that label.",
                    "Providers and reviewers typically evaluate actual behavior, not labels alone.",
                ],
            },
            {
                "title": "7. Provider-specific commitments",
                "bullets": [
                    "Comply with Spotify Developer Terms and Spotify Developer Policy.",
                    "Comply with the Google API Services User Data Policy.",
                    "Comply with the YouTube API Services Terms and Developer Policies.",
                    "Request minimum necessary permissions and accurately explain user-facing value.",
                    "Avoid deceptive or unauthorized use of provider APIs.",
                ],
            },
        ],
    },
]

SOURCES = [
    "Spotify Developer Policy — https://developer.spotify.com/policy",
    "Spotify Developer Terms — https://developer.spotify.com/terms",
    "Spotify Web API rate limits — https://developer.spotify.com/documentation/web-api/concepts/rate-limits",
    "Spotify quota modes — https://developer.spotify.com/documentation/web-api/concepts/quota-modes",
    "Google API Services User Data Policy — https://developers.google.com/terms/api-services-user-data-policy",
    "YouTube API Services Developer Policies — https://developers.google.com/youtube/terms/developer-policies",
    "YouTube quota and compliance audits — https://developers.google.com/youtube/v3/guides/quota_and_compliance_audits",
]


def set_cell_fill(cell, fill: str, border: str | None = None) -> None:
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))
    if border:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge in ("top", "left", "bottom", "right"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:color"), border)
            tc_borders.append(element)


def apply_paragraph_style(paragraph, *, color: RGBColor = TEXT, size: int = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def set_cell_padding(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for side, amount in values.items():
        side_element = tc_mar.find(qn(f"w:{side}"))
        if side_element is None:
            side_element = OxmlElement(f"w:{side}")
            tc_mar.append(side_element)
        side_element.set(qn("w:w"), str(amount))
        side_element.set(qn("w:type"), "dxa")


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
    brand_run = brand.add_run("SongZip")
    brand_run.font.name = "Aptos Display"
    brand_run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    brand_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    brand_run.font.size = Pt(11)
    brand_run.font.bold = True
    brand_run.font.color.rgb = ACCENT
    brand.paragraph_format.space_after = Pt(8)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(title)
    title_run.font.name = "Aptos Display"
    title_run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    title_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = TEXT
    title_paragraph.paragraph_format.space_after = Pt(4)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(f"{subtitle} • Last updated May 16, 2026")
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = TEXT_SOFT
    subtitle_paragraph.paragraph_format.space_after = Pt(12)


def add_summary_box(doc: Document, summary: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, BOX_FILL, BOX_EDGE)
    set_cell_padding(cell, top=120, start=140, bottom=120, end=140)

    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run("Public summary")
    title_run.font.size = Pt(9.5)
    title_run.font.bold = True
    title_run.font.color.rgb = ACCENT
    title_paragraph.paragraph_format.space_after = Pt(4)

    body = cell.add_paragraph(summary)
    body.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(body, color=TEXT, size=11)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_section(doc: Document, title: str, paragraphs: list[str] | None = None, bullets: list[str] | None = None) -> None:
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.font.name = "Aptos Display"
    heading_run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    heading_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.color.rgb = TEXT
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)

    for text in paragraphs or []:
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.12
        apply_paragraph_style(paragraph, color=TEXT, size=11)

    for item in bullets or []:
        bullet = doc.add_paragraph()
        bullet.add_run("• ")
        bullet.add_run(item)
        bullet.paragraph_format.space_after = Pt(4)
        bullet.paragraph_format.left_indent = Inches(0.18)
        bullet.paragraph_format.first_line_indent = Inches(-0.14)
        bullet.paragraph_format.line_spacing = 1.08
        apply_paragraph_style(bullet, color=TEXT, size=11)


def add_sources(doc: Document) -> None:
    add_section(doc, "Source references")
    for source in SOURCES:
        bullet = doc.add_paragraph()
        bullet.add_run("• ")
        bullet.add_run(source)
        bullet.paragraph_format.space_after = Pt(4)
        bullet.paragraph_format.left_indent = Inches(0.18)
        bullet.paragraph_format.first_line_indent = Inches(-0.14)
        apply_paragraph_style(bullet, color=TEXT, size=10.5)


def set_document_defaults(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run(title)
    footer_run.font.name = "Aptos"
    footer_run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    footer_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = TEXT_SOFT


def build_document(policy: dict, output_dir: Path) -> Path:
    doc = Document()
    set_document_defaults(doc, policy["title"])
    add_title_block(doc, policy["title"], policy["subtitle"])
    add_summary_box(doc, policy["summary"])

    note = doc.add_paragraph(
        "This document is a transparency draft for the current SongZip dashboard build. "
        "It is not legal advice, and the product's real behavior controls."
    )
    note.paragraph_format.space_after = Pt(10)
    apply_paragraph_style(note, color=TEXT_SOFT, size=10.5)

    for section in policy["sections"]:
        add_section(
            doc,
            section["title"],
            paragraphs=section.get("paragraphs"),
            bullets=section.get("bullets"),
        )

    add_sources(doc)

    output_path = output_dir / policy["output"]
    doc.save(output_path)
    return output_path


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    output_dir = repo_root / "local-web-ui"
    output_dir.mkdir(parents=True, exist_ok=True)

    for policy in POLICIES:
        output_path = build_document(policy, output_dir)
        print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
